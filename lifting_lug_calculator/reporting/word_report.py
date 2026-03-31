"""Word report generation."""

import io
from dataclasses import dataclass

from docx import Document
from docx.oxml.ns import qn

from lifting_lug_calculator.core.constants import ALLOWABLE_RULES


@dataclass(frozen=True)
class ReportParams:
    fv_ton: float
    fv: float
    kd: float
    theta: float
    t: float
    tp: float
    d: float
    dp: float
    r: float
    h: float
    w: float
    hf: float
    fy: float
    fu: float
    ns: float


def _add_metric_section(doc: Document, heading: str, description: str, equation: str, allowable: str, passed: bool) -> None:
    doc.add_heading(heading, level=2)
    doc.add_paragraph(description)
    doc.add_paragraph(equation)
    doc.add_paragraph(allowable)
    doc.add_paragraph("结论：" + ("满足要求" if passed else "不满足要求"))


def generate_word_report(params: ReportParams, results: dict) -> io.BytesIO:
    """Generate the Word-format report."""
    doc = Document()
    doc.styles["Normal"].font.name = "宋体"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.add_heading("吊耳应力校验计算书", 0)

    doc.add_heading("一、基本参数", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "参数"
    hdr_cells[1].text = "符号"
    hdr_cells[2].text = "数值"
    hdr_cells[3].text = "单位"

    param_data = [
        ("单只吊耳受力", "Fv", f"{params.fv_ton} t ({params.fv:.2f} kN)", "—"),
        ("动载系数", "Kd", str(params.kd), "—"),
        ("吊索夹角", "θ", str(params.theta), "°"),
        ("材料屈服强度", "Fy", str(params.fy), "MPa"),
        ("材料抗拉强度", "Fu", str(params.fu), "MPa"),
        ("结构安全系数", "Ns", str(params.ns), "—"),
        ("吊耳主板厚", "t", str(params.t), "mm"),
        ("单侧加强板厚", "tp", str(params.tp), "mm"),
        ("有效厚度", "t_eff", str(results["t_eff"]), "mm"),
        ("销轴孔孔径", "d", str(params.d), "mm"),
        ("销轴实际直径", "dp", str(params.dp), "mm"),
        ("孔心至外边缘距离", "R", str(params.r), "mm"),
        ("圆心到底边距离", "H", str(params.h), "mm"),
        ("吊耳根部宽度", "W", str(params.w), "mm"),
        ("焊缝角边尺寸", "hf", str(params.hf), "mm"),
    ]

    for item in param_data:
        row_cells = table.add_row().cells
        for index in range(4):
            row_cells[index].text = item[index]

    doc.add_heading("二、校验结果汇总", level=1)
    res_table = doc.add_table(rows=1, cols=4)
    res_table.style = "Table Grid"
    hdr_cells = res_table.rows[0].cells
    hdr_cells[0].text = "校验项"
    hdr_cells[1].text = "计算应力 (MPa)"
    hdr_cells[2].text = "许用应力 (MPa)"
    hdr_cells[3].text = "结论"

    metrics = [
        ("1. 拉应力 (Tensile)", results["metrics"]["tensile"]["value"], results["metrics"]["tensile"]["allowable"], results["metrics"]["tensile"]["passed"]),
        ("2. 撕裂应力 (Tear-out)", results["metrics"]["tear_out"]["value"], results["metrics"]["tear_out"]["allowable"], results["metrics"]["tear_out"]["passed"]),
        ("3. 承压应力 (Bearing)", results["metrics"]["bearing"]["value"], results["metrics"]["bearing"]["allowable"], results["metrics"]["bearing"]["passed"]),
        ("4. 焊缝应力 (Weld)", results["metrics"]["weld"]["value"], results["metrics"]["weld"]["allowable"], results["metrics"]["weld"]["passed"]),
    ]

    for name, calc_val, allow_val, is_pass in metrics:
        row = res_table.add_row().cells
        row[0].text = name
        row[1].text = f"{calc_val:.2f}" if calc_val is not None else "N/A"
        row[2].text = f"{allow_val:.2f}"
        row[3].text = "通过" if is_pass else "不通过"

    doc.add_heading("三、详细计算过程", level=1)
    doc.add_heading("3.1 换算设计拉力 P", level=2)
    doc.add_paragraph(
        f"P = (Fv × 1000 × Kd) / sin(θ) = ({params.fv:.2f} × 1000 × {params.kd}) / sin({params.theta}°) = {results['p_design']:.2f} N"
    )

    _add_metric_section(
        doc,
        heading="3.2 销轴孔净截面拉应力 σ_t",
        description="计算孔横截面处的拉伸应力：\nσ_t = P / [2 × t_eff × (R - d/2)]",
        equation=f"σ_t = {results['p_design']:.2f} / [2 × {results['t_eff']} × ({params.r} - {params.d}/2)] = {results['metrics']['tensile']['value']:.2f} MPa",
        allowable=f"许用拉应力 {ALLOWABLE_RULES['tensile']['label']} = {results['metrics']['tensile']['allowable']:.2f} MPa",
        passed=results["metrics"]["tensile"]["passed"],
    )
    _add_metric_section(
        doc,
        heading="3.3 销轴孔端部撕裂应力 τ_t",
        description="计算孔端部剪切撕裂面的切应力：\nτ_t = P / [2 × t_eff × (R - d/2)]",
        equation=f"τ_t = {results['p_design']:.2f} / [2 × {results['t_eff']} × ({params.r} - {params.d}/2)] = {results['metrics']['tear_out']['value']:.2f} MPa",
        allowable=f"许用剪应力 {ALLOWABLE_RULES['tear_out']['label']} = {results['metrics']['tear_out']['allowable']:.2f} MPa",
        passed=results["metrics"]["tear_out"]["passed"],
    )
    _add_metric_section(
        doc,
        heading="3.4 销轴与孔壁局部承压应力 σ_p",
        description="计算孔壁接触面的局部挤压应力：\nσ_p = P / (t_eff × dp)",
        equation=f"σ_p = {results['p_design']:.2f} / ({results['t_eff']} × {params.dp}) = {results['metrics']['bearing']['value']:.2f} MPa",
        allowable=f"许用挤压应力 {ALLOWABLE_RULES['bearing']['label']} = {results['metrics']['bearing']['allowable']:.2f} MPa",
        passed=results["metrics"]["bearing"]["passed"],
    )

    doc.add_heading("3.5 吊耳根部角焊缝剪应力 τ_w", level=2)
    doc.add_paragraph("计算主板与母材连接焊缝的平均切应力（保守计算，仅计算两侧角焊缝）：\nτ_w = P / (2 × 0.7 × hf × W)")
    doc.add_paragraph("注：当存在加强板时，计算仍偏向于偏保守的假设——所有拉力由主体板周围的焊缝传递。")
    doc.add_paragraph(f"τ_w = {results['p_design']:.2f} / (2 × 0.7 × {params.hf} × {params.w}) = {results['metrics']['weld']['value']:.2f} MPa")
    doc.add_paragraph(f"许用焊缝应力 {ALLOWABLE_RULES['weld']['label']} = {results['metrics']['weld']['allowable']:.2f} MPa")
    doc.add_paragraph("结论：" + ("满足要求" if results["metrics"]["weld"]["passed"] else "不满足要求"))

    doc.add_heading("四、最终结论", level=1)
    if all(results["metrics"][name]["passed"] for name in ("tensile", "tear_out", "bearing", "weld")):
        doc.add_paragraph("该吊耳设计满足所有受力校验要求，设计安全合理。")
    else:
        doc.add_paragraph("该吊耳设计存在风险，部分指标超出许用应力，请修改参数！")

    byte_io = io.BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io
