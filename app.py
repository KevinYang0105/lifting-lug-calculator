"""
吊耳计算器 (Lifting Lug Calculator)
基于工程力学公式，计算并校验吊耳的四项核心应力指标。
"""

import io
import math
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from docx import Document
from docx.oxml.ns import qn

# ─────────────────────────────────────────────
#  页面基础配置
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="吊耳计算器 | Lifting Lug Calculator",
    page_icon="🔩",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────
#  自定义 CSS（工程风格深色主题）
# ─────────────────────────────────────────────
st.markdown("""
<style>
/* 全局字体 */
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
}

/* 主背景 */
.stApp {
    background: linear-gradient(135deg, #0f1117 0%, #1a1d27 50%, #0f1117 100%);
    color: #e8eaf0;
}

/* 侧边栏 */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #161b2e 0%, #1e2235 100%);
    border-right: 1px solid rgba(99, 179, 237, 0.15);
}

/* 标题样式 */
.main-title {
    font-size: 2rem;
    font-weight: 700;
    background: linear-gradient(90deg, #63b3ed, #76e4f7, #9f7aea);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    margin-bottom: 0.25rem;
}

.subtitle {
    color: #718096;
    font-size: 0.9rem;
    margin-bottom: 2rem;
    font-weight: 400;
    letter-spacing: 0.05em;
}

/* 状态卡片 */
.status-card {
    background: rgba(26, 32, 55, 0.8);
    border: 1px solid rgba(99, 179, 237, 0.1);
    border-radius: 12px;
    padding: 1.2rem 1.4rem;
    margin-bottom: 0.75rem;
    backdrop-filter: blur(10px);
    transition: all 0.2s ease;
}

.status-card:hover {
    border-color: rgba(99, 179, 237, 0.3);
    transform: translateY(-1px);
}

.status-pass {
    border-left: 4px solid #48bb78;
    background: rgba(72, 187, 120, 0.05);
}

.status-fail {
    border-left: 4px solid #fc8181;
    background: rgba(252, 129, 129, 0.05);
}

.stress-name {
    font-size: 0.85rem;
    color: #a0aec0;
    font-weight: 500;
    letter-spacing: 0.03em;
    margin-bottom: 0.5rem;
}

.stress-values {
    display: flex;
    align-items: center;
    gap: 1rem;
}

.stress-actual {
    font-size: 1.6rem;
    font-weight: 700;
    font-family: 'JetBrains Mono', monospace;
    color: #e2e8f0;
}

.stress-allowable {
    font-size: 0.8rem;
    color: #718096;
}

.status-badge-pass {
    background: rgba(72, 187, 120, 0.15);
    color: #68d391;
    border: 1px solid rgba(72, 187, 120, 0.3);
    border-radius: 20px;
    padding: 0.2rem 0.8rem;
    font-size: 0.75rem;
    font-weight: 600;
    letter-spacing: 0.05em;
}

.status-badge-fail {
    background: rgba(252, 129, 129, 0.15);
    color: #fc8181;
    border: 1px solid rgba(252, 129, 129, 0.3);
    border-radius: 20px;
    padding: 0.2rem 0.8rem;
    font-size: 0.75rem;
    font-weight: 600;
    letter-spacing: 0.05em;
}

/* 分割线 */
.section-divider {
    border: none;
    border-top: 1px solid rgba(99, 179, 237, 0.1);
    margin: 1.5rem 0;
}

/* 计算书区域 */
.calc-report {
    background: rgba(15, 17, 28, 0.7);
    border: 1px solid rgba(99, 179, 237, 0.1);
    border-radius: 12px;
    padding: 1.5rem 2rem;
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.85rem;
}

/* Streamlit 组件微调 */

/* 加亮侧边栏输入框上方的文本说明（Label） */
div[data-testid="stWidgetLabel"] p, 
div[data-testid="stWidgetLabel"] label,
label span,
.stCheckbox p {
    color: #f7fafc !important;
    font-weight: 500 !important;
}

[data-testid="stMetric"] {
    background: rgba(26,32,55,0.5);
    border-radius: 8px;
    padding: 0.5rem;
}

/* 加亮主界面 st.metric 的标签文字和图标 */
[data-testid="stMetricLabel"] p {
    color: #f7fafc !important;
    font-weight: 500 !important;
}
[data-testid="stMetricLabel"] svg {
    fill: #f7fafc !important;
    color: #f7fafc !important;
}

/* 加亮 st.metric 的具体输出数值结果 */
[data-testid="stMetricValue"] div {
    color: #ffffff !important;
    font-weight: 700 !important;
}

div.stButton > button {
    background: linear-gradient(135deg, #3182ce, #2b6cb0);
    color: white;
    border: none;
    border-radius: 8px;
    padding: 0.5rem 1.5rem;
    font-weight: 600;
    width: 100%;
    transition: all 0.2s;
}

div.stButton > button:hover {
    background: linear-gradient(135deg, #4299e1, #3182ce);
    transform: translateY(-1px);
    box-shadow: 0 4px 15px rgba(49,130,206,0.4);
}

.sidebar-section-title {
    font-size: 0.7rem;
    font-weight: 600;
    letter-spacing: 0.12em;
    text-transform: uppercase;
    color: #4299e1;
    margin: 1.2rem 0 0.5rem 0;
}

/* 导出Word按钮高亮外观 (淡蓝色背景 + 醒目深色文字) */
[data-testid="stDownloadButton"] button {
    background: linear-gradient(135deg, #e0f2fe, #7dd3fc) !important;
    color: #0c4a6e !important;
    font-weight: 800 !important;
    font-size: 1rem !important;
    border: 2px solid #38bdf8 !important;
    border-radius: 8px !important;
    box-shadow: 0 4px 10px rgba(56, 189, 248, 0.4) !important;
    transition: all 0.2s ease !important;
}

[data-testid="stDownloadButton"] button:hover {
    background: linear-gradient(135deg, #bae6fd, #38bdf8) !important;
    transform: translateY(-2px) !important;
    box-shadow: 0 6px 15px rgba(56, 189, 248, 0.6) !important;
}

/* 展开计算过程区域（Expander）标题高亮外观 */
[data-testid="stExpander"] summary {
    background: linear-gradient(135deg, #e0f2fe, #7dd3fc) !important;
    border-radius: 8px !important;
    border: 1px solid #38bdf8 !important;
    padding: 0.5rem 1rem !important;
    margin-bottom: 0.5rem !important;
}

[data-testid="stExpander"] summary p {
    color: #0c4a6e !important;
    font-weight: 800 !important;
    font-size: 1.05rem !important;
}

[data-testid="stExpander"] summary svg {
    color: #0c4a6e !important;
    fill: #0c4a6e !important;
}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
#  材质数据库
# ─────────────────────────────────────────────
MATERIAL_DB = {
    "Q355B": {"fy": 355, "fu": 470},
    "Q235B": {"fy": 235, "fu": 370},
    "Q420B": {"fy": 420, "fu": 520},
    "SS400":  {"fy": 245, "fu": 400},
    "自定义": {"fy": 355, "fu": 470},
}

ALLOWABLE_RULES = {
    "tensile": {
        "label": "[σ_t] = Fy / Ns",
        "compute": lambda fy, fu, ns: fy / ns,
    },
    "tear_out": {
        "label": "[τ_v] = 0.58 × Fy / Ns",
        "compute": lambda fy, fu, ns: (0.58 * fy) / ns,
    },
    "bearing": {
        "label": "[σ_b] = 1.25 × Fy / Ns",
        "compute": lambda fy, fu, ns: (1.25 * fy) / ns,
    },
    "weld": {
        "label": "[τ_w] = 0.3 × Fu",
        "compute": lambda fy, fu, ns: 0.3 * fu,
    },
}

# ─────────────────────────────────────────────
#  核心计算函数 (统一接口)
# ─────────────────────────────────────────────
def calculate_lug_stresses(fv, kd, theta_deg, t, t_p, d, dp, r, w, hf, fy, fu, ns):
    """
    计算吊耳各项应力并进行安全校验

    参数:
    fv: 单只吊耳垂直基础载荷 (kN)
    kd: 动载系数
    theta_deg: 吊索与水平面夹角 (度)
    t: 吊耳主板厚度 (mm)
    t_p: 单侧加强板厚度 (mm)，无加强板时取 0
    d: 销轴孔孔径 (mm)
    dp: 销轴实际直径 (mm)
    r: 销轴孔中心至外边缘距离 (mm)
    w: 吊耳根部宽度 (mm)
    hf: 角焊缝焊脚尺寸 (mm)
    fy: 材料屈服强度 (MPa)
    fu: 材料抗拉强度 (MPa)
    ns: 安全系数

    返回:
    dict: 包含各项应力计算结果及是否通过校验的布尔值
    """
    try:
        # 1. 基础载荷换算
        theta_rad = math.radians(theta_deg)
        p_design = (fv * 1000 * kd) / math.sin(theta_rad)

        # 2. 计算有效厚度 (主板 + 双侧加强板)
        t_eff = t + 2 * t_p

        # 几何校验
        if r <= d / 2 or t_eff <= 0 or dp <= 0 or w <= 0 or hf <= 0:
            return {"error": "几何尺寸输入错误（如 R 必须大于 d/2，厚度/直径不能为 0）"}
        if dp > d:
            return {"error": "销轴直径 dp 不能大于孔径 d。"}

        # 3. 核心应力计算
        # --- 中间面积值 (用于计算书展示) ---
        net_area = 2 * t_eff * (r - d / 2)
        shear_area = 2 * t_eff * (r - d / 2)
        bearing_area = t_eff * dp
        weld_throat = 0.707 * hf
        weld_area = weld_throat * 2 * w

        # 指标一：销轴孔净截面拉应力
        tensile_stress = p_design / net_area
        allowable_tensile = ALLOWABLE_RULES["tensile"]["compute"](fy, fu, ns)

        # 指标二：孔端面撕裂应力 (双面剪切)
        shear_stress = p_design / shear_area
        allowable_shear = ALLOWABLE_RULES["tear_out"]["compute"](fy, fu, ns)

        # 指标三：销轴孔局部承压应力
        bearing_stress = p_design / bearing_area
        allowable_bearing = ALLOWABLE_RULES["bearing"]["compute"](fy, fu, ns)

        # 指标四：底部焊缝综合剪应力 (仅通过主板传递)
        weld_shear_stress = p_design / weld_area
        allowable_weld = ALLOWABLE_RULES["weld"]["compute"](fy, fu, ns)

        # 4. 组装返回数据
        results = {
            "p_design": p_design,
            "t_eff": t_eff,
            "areas": {
                "net_area": net_area,
                "shear_area": shear_area,
                "bearing_area": bearing_area,
                "weld_throat": weld_throat,
                "weld_area": weld_area,
            },
            "metrics": {
                "tensile": {
                    "value": tensile_stress,
                    "allowable": allowable_tensile,
                    "passed": tensile_stress <= allowable_tensile,
                },
                "tear_out": {
                    "value": shear_stress,
                    "allowable": allowable_shear,
                    "passed": shear_stress <= allowable_shear,
                },
                "bearing": {
                    "value": bearing_stress,
                    "allowable": allowable_bearing,
                    "passed": bearing_stress <= allowable_bearing,
                },
                "weld": {
                    "value": weld_shear_stress,
                    "allowable": allowable_weld,
                    "passed": weld_shear_stress <= allowable_weld,
                },
            },
        }
        return results

    except Exception as e:
        return {"error": f"计算过程发生异常: {str(e)}"}

def generate_word_report(params, results):
    """生成 Word 格式的计算书"""
    doc = Document()
    # 设置全文字体支持中文
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    
    doc.add_heading('吊耳应力校验计算书', 0)
    
    doc.add_heading('一、基本参数', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '参数'
    hdr_cells[1].text = '符号'
    hdr_cells[2].text = '数值'
    hdr_cells[3].text = '单位'
    
    param_data = [
        ('单只吊耳受力', 'Fv', f"{params['fv_ton']} t ({params['fv']:.2f} kN)", '—'),
        ('动载系数', 'Kd', str(params['kd']), '—'),
        ('吊索夹角', 'θ', str(params['theta']), '°'),
        ('材料屈服强度', 'Fy', str(params['fy']), 'MPa'),
        ('材料抗拉强度', 'Fu', str(params['fu']), 'MPa'),
        ('结构安全系数', 'Ns', str(params['ns']), '—'),
        ('吊耳主板厚', 't', str(params['t']), 'mm'),
        ('单侧加强板厚', 'tp', str(params['tp']), 'mm'),
        ('有效厚度', 't_eff', str(results['t_eff']), 'mm'),
        ('销轴孔孔径', 'd', str(params['d']), 'mm'),
        ('销轴实际直径', 'dp', str(params['dp']), 'mm'),
        ('孔心至外边缘距离', 'R', str(params['r']), 'mm'),
        ('吊耳根部宽度', 'W', str(params['w']), 'mm'),
        ('焊缝角边尺寸', 'hf', str(params['hf']), 'mm'),
    ]
    
    for item in param_data:
        row_cells = table.add_row().cells
        for i in range(4):
            row_cells[i].text = item[i]
            
    doc.add_heading('二、校验结果汇总', level=1)
    res_table = doc.add_table(rows=1, cols=4)
    res_table.style = 'Table Grid'
    hdr_cells = res_table.rows[0].cells
    hdr_cells[0].text = '校验项'
    hdr_cells[1].text = '计算应力 (MPa)'
    hdr_cells[2].text = '许用应力 (MPa)'
    hdr_cells[3].text = '结论'
    
    metrics = [
        ('1. 拉应力 (Tensile)', results['metrics']['tensile']['value'], results['metrics']['tensile']['allowable'], results['metrics']['tensile']['passed']),
        ('2. 撕裂应力 (Tear-out)', results['metrics']['tear_out']['value'], results['metrics']['tear_out']['allowable'], results['metrics']['tear_out']['passed']),
        ('3. 承压应力 (Bearing)', results['metrics']['bearing']['value'], results['metrics']['bearing']['allowable'], results['metrics']['bearing']['passed']),
        ('4. 焊缝应力 (Weld)', results['metrics']['weld']['value'], results['metrics']['weld']['allowable'], results['metrics']['weld']['passed']),
    ]
    
    for name, calc_val, allow_val, is_pass in metrics:
        row = res_table.add_row().cells
        row[0].text = name
        row[1].text = f"{calc_val:.2f}" if calc_val is not None else "N/A"
        row[2].text = f"{allow_val:.2f}"
        row[3].text = "通过" if is_pass else "不通过"
        
    doc.add_heading('三、详细计算过程', level=1)
    
    # 设计拉力
    doc.add_heading('3.1 换算设计拉力 P', level=2)
    doc.add_paragraph(f"P = (Fv × 1000 × Kd) / sin(θ) = ({params['fv']:.2f} × 1000 × {params['kd']}) / sin({params['theta']}°) = {results['p_design']:.2f} N")
    
    # 拉应力
    doc.add_heading('3.2 销轴孔净截面拉应力 σ_t', level=2)
    doc.add_paragraph("计算孔横截面处的拉伸应力：\nσ_t = P / [2 × t_eff × (R - d/2)]")
    val = results['metrics']['tensile']['value']
    if val is not None:
        doc.add_paragraph(f"σ_t = {results['p_design']:.2f} / [2 × {results['t_eff']} × ({params['r']} - {params['d']}/2)] = {val:.2f} MPa")
        doc.add_paragraph(
            f"许用拉应力 {ALLOWABLE_RULES['tensile']['label']} = {results['metrics']['tensile']['allowable']:.2f} MPa"
        )
        doc.add_paragraph("结论：" + ("满足要求" if results['metrics']['tensile']['passed'] else "不满足要求"))
    else:
        doc.add_paragraph("尺寸异常无法计算。")
        
    # 撕裂应力
    doc.add_heading('3.3 销轴孔端部撕裂应力 τ_t', level=2)
    doc.add_paragraph("计算孔端部剪切撕裂面的切应力：\nτ_t = P / [2 × t_eff × (R - d/2)]")
    val = results['metrics']['tear_out']['value']
    if val is not None:
        doc.add_paragraph(f"τ_t = {results['p_design']:.2f} / [2 × {results['t_eff']} × ({params['r']} - {params['d']}/2)] = {val:.2f} MPa")
        doc.add_paragraph(
            f"许用剪应力 {ALLOWABLE_RULES['tear_out']['label']} = {results['metrics']['tear_out']['allowable']:.2f} MPa"
        )
        doc.add_paragraph("结论：" + ("满足要求" if results['metrics']['tear_out']['passed'] else "不满足要求"))
    else:
        doc.add_paragraph("尺寸异常无法计算。")
        
    # 承压应力
    doc.add_heading('3.4 销轴与孔壁局部承压应力 σ_p', level=2)
    doc.add_paragraph("计算孔壁接触面的局部挤压应力：\nσ_p = P / (t_eff × dp)")
    val = results['metrics']['bearing']['value']
    if val is not None:
        doc.add_paragraph(f"σ_p = {results['p_design']:.2f} / ({results['t_eff']} × {params['dp']}) = {val:.2f} MPa")
        doc.add_paragraph(
            f"许用挤压应力 {ALLOWABLE_RULES['bearing']['label']} = {results['metrics']['bearing']['allowable']:.2f} MPa"
        )
        doc.add_paragraph("结论：" + ("满足要求" if results['metrics']['bearing']['passed'] else "不满足要求"))
    else:
        doc.add_paragraph("尺寸异常无法计算。")
        
    # 焊缝应力
    doc.add_heading('3.5 吊耳根部角焊缝剪应力 τ_w', level=2)
    doc.add_paragraph("计算主板与母材连接焊缝的平均切应力（保守计算，仅计算两侧角焊缝）：\nτ_w = P / (2 × 0.7 × hf × W)")
    doc.add_paragraph("注：当存在加强板时，计算仍偏向于偏保守的假设——所有拉力由主体板周围的焊缝传递。")
    val = results['metrics']['weld']['value']
    if val is not None:
        doc.add_paragraph(f"τ_w = {results['p_design']:.2f} / (2 × 0.7 × {params['hf']} × {params['w']}) = {val:.2f} MPa")
        doc.add_paragraph(
            f"许用焊缝应力 {ALLOWABLE_RULES['weld']['label']} = {results['metrics']['weld']['allowable']:.2f} MPa"
        )
        doc.add_paragraph("结论：" + ("满足要求" if results['metrics']['weld']['passed'] else "不满足要求"))
    else:
        doc.add_paragraph("尺寸异常无法计算。")
        
    doc.add_heading('四、最终结论', level=1)
    if all([results['metrics']['tensile']['passed'], 
            results['metrics']['tear_out']['passed'], 
            results['metrics']['bearing']['passed'], 
            results['metrics']['weld']['passed']]):
        doc.add_paragraph('该吊耳设计满足所有受力校验要求，设计安全合理。')
    else:
        doc.add_paragraph('该吊耳设计存在风险，部分指标超出许用应力，请修改参数！')
        
    byte_io = io.BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io


def draw_lug_diagram(W, R, d, t, t_p, theta):
    """绘制标准正交投影双视图（正视图 + 侧视图）"""
    # 设置画布
    fig, (ax1, ax2) = plt.subplots(
        1, 2, figsize=(10, 6),
        gridspec_kw={'width_ratios': [2.2, 1]},
    )
    bg_color = '#1e2235'
    fig.patch.set_facecolor(bg_color)

    # 通用颜色
    edge_color = '#63b3ed'
    fill_color = '#2d3748'
    reinf_edge = '#7b93db'
    text_color = '#e2e8f0'
    dim_color  = '#a0aec0'
    hole_fill  = bg_color
    centerline = '#718096'

    hole_r = d / 2
    main_height = 80  # 矩形主体高度
    top_y = main_height  # 孔心 y 坐标
    head_radius = R
    # 加强板外半径（取孔径外侧一圈，约 0.7R）
    Dp_R = min(R * 0.7, R - 2) if t_p > 0 else 0

    # ====================
    # 左侧：正视图 (Front View)
    # ====================
    ax1.set_aspect('equal')
    ax1.axis('off')
    ax1.set_facecolor(bg_color)

    # 主体轮廓
    ax1.add_patch(patches.Rectangle(
        (-W/2, 0), W, main_height,
        linewidth=2, edgecolor=edge_color, facecolor=fill_color, alpha=0.8,
    ))
    ax1.add_patch(patches.Arc(
        (0, top_y), 2 * head_radius, 2 * head_radius, theta1=0, theta2=180,
        linewidth=2, edgecolor=edge_color,
    ))
    if W / 2 > head_radius:
        ax1.plot([-W / 2, -head_radius], [top_y, top_y], color=edge_color, lw=2)
        ax1.plot([head_radius, W / 2], [top_y, top_y], color=edge_color, lw=2)

    # 加强板轮廓（虚线圆）
    if t_p > 0:
        ax1.add_patch(patches.Circle(
            (0, top_y), Dp_R,
            linewidth=1.5, edgecolor=reinf_edge, linestyle='--', facecolor='none',
        ))

    # 销轴孔
    ax1.add_patch(patches.Circle(
        (0, top_y), hole_r,
        linewidth=2, edgecolor=edge_color, facecolor=hole_fill,
    ))

    # 中心线
    half_span = max(W / 2, head_radius)
    ax1.plot([-half_span - 12, half_span + 12], [top_y, top_y], color=centerline, linestyle='-.', lw=0.8)
    ax1.plot([0, 0], [-10, top_y + head_radius + 12], color=centerline, linestyle='-.', lw=0.8)

    # 标注尺寸
    # 孔径 d
    ax1.annotate('', xy=(-hole_r, top_y), xytext=(hole_r, top_y),
                 arrowprops=dict(arrowstyle='<->', color=edge_color, lw=1.2))
    ax1.text(0, top_y + hole_r + 4, f'd={d:.0f}', ha='center',
             color=edge_color, fontsize=8, fontweight='bold')

    # 半径 R
    angle_45 = math.radians(45)
    ax1.annotate('', xy=(0, top_y),
                 xytext=(head_radius * math.cos(angle_45), top_y + head_radius * math.sin(angle_45)),
                 arrowprops=dict(arrowstyle='<-', color=dim_color, lw=1))
    ax1.text(head_radius / 2 * math.cos(angle_45) + 3, top_y + head_radius / 2 * math.sin(angle_45) + 3,
             f'R={R:.0f}', color=text_color, fontsize=8)

    # 宽度 W
    ax1.annotate('', xy=(-W/2, -12), xytext=(W/2, -12),
                 arrowprops=dict(arrowstyle='<->', color=dim_color, lw=1))
    ax1.text(0, -22, f'W={W:.0f}', ha='center', color=text_color, fontsize=8)

    # 载荷 P 矢量
    force_len = 35
    dx = force_len * math.cos(math.radians(theta))
    dy = force_len * math.sin(math.radians(theta))
    ax1.annotate('', xy=(0, top_y), xytext=(dx, top_y + dy),
                 arrowprops=dict(arrowstyle='<-', color='#fc8181', lw=2))
    ax1.text(dx + 2, top_y + dy + 4, 'P', color='#fc8181', fontsize=10, fontweight='bold')
    # 角度参考线
    ax1.plot([0, force_len * 1.1], [top_y, top_y], color=centerline, linestyle=':', lw=0.8)
    # 角度弧
    arc_r = force_len * 0.4
    theta_arc = patches.Arc((0, top_y), arc_r*2, arc_r*2,
                            theta1=0, theta2=theta,
                            linewidth=1, edgecolor='#fc8181', linestyle='-')
    ax1.add_patch(theta_arc)
    ax1.text(arc_r * 0.6, top_y + arc_r * 0.3, f'θ={theta}°',
             color='#fc8181', fontsize=7)

    ax1.text(0, -35, 'Front View', ha='center', color=text_color, fontsize=9, fontstyle='italic')

    ax1.set_xlim(-half_span - 25, half_span + 45)
    ax1.set_ylim(-42, top_y + head_radius + 20)

    # ====================
    # 右侧：侧视图 (Side View)
    # ====================
    ax2.set_aspect('equal')
    ax2.axis('off')
    ax2.set_facecolor(bg_color)

    total_height = main_height + head_radius  # 包含半圆头的总高度

    # 主板
    ax2.add_patch(patches.Rectangle(
        (-t/2, 0), t, total_height,
        linewidth=2, edgecolor=edge_color, facecolor=fill_color, alpha=0.8,
    ))

    # 加强板
    if t_p > 0:
        reinf_h = 2 * Dp_R
        reinf_y = top_y - Dp_R
        # 左侧
        ax2.add_patch(patches.Rectangle(
            (-t/2 - t_p, reinf_y), t_p, reinf_h,
            linewidth=1.5, edgecolor=reinf_edge, facecolor='#2a2f50', alpha=0.7,
        ))
        # 右侧
        ax2.add_patch(patches.Rectangle(
            (t/2, reinf_y), t_p, reinf_h,
            linewidth=1.5, edgecolor=reinf_edge, facecolor='#2a2f50', alpha=0.7,
        ))
        # tp 标注（右侧加强板）
        label_y = reinf_y + reinf_h + 6
        ax2.annotate('', xy=(t/2, label_y), xytext=(t/2 + t_p, label_y),
                     arrowprops=dict(arrowstyle='<->', color=reinf_edge, lw=1))
        ax2.text(t/2 + t_p/2, label_y + 4, f'tp={t_p:.0f}',
             ha='center', color=reinf_edge, fontsize=7, fontweight='bold')

    # 销轴孔贯穿虚线
    side_extent = t/2 + (t_p if t_p > 0 else 0)
    ax2.plot([-side_extent - 5, side_extent + 5],
             [top_y + hole_r, top_y + hole_r], color=centerline, linestyle='--', lw=0.8)
    ax2.plot([-side_extent - 5, side_extent + 5],
             [top_y - hole_r, top_y - hole_r], color=centerline, linestyle='--', lw=0.8)

    # 中心线
    ax2.plot([0, 0], [-10, total_height + 12], color=centerline, linestyle='-.', lw=0.8)
    ax2.plot([-side_extent - 15, side_extent + 15],
             [top_y, top_y], color=centerline, linestyle='-.', lw=0.8)

    # t 标注
    t_label_y = 25
    ax2.annotate('', xy=(-t/2, t_label_y), xytext=(t/2, t_label_y),
                 arrowprops=dict(arrowstyle='<->', color=dim_color, lw=1))
    ax2.text(0, t_label_y + 4, f't={t:.0f}', ha='center',
             color=text_color, fontsize=8)

    ax2.text(0, -35, 'Side View', ha='center', color=text_color, fontsize=9, fontstyle='italic')

    ax2.set_xlim(-side_extent - 25, side_extent + 25)
    ax2.set_ylim(-42, total_height + 20)

    plt.tight_layout(pad=1.0)
    return fig


# ─────────────────────────────────────────────
#  侧边栏 — 输入参数区
# ─────────────────────────────────────────────
with st.sidebar:
    st.markdown('<p class="main-title" style="font-size:1.3rem;">🔩 吊耳计算器</p>', unsafe_allow_html=True)
    st.markdown('<p style="color:#718096; font-size:0.78rem;">Lifting Lug Stress Calculator</p>', unsafe_allow_html=True)

    st.divider()

    # —— 材质参数 ——
    st.markdown('<p class="sidebar-section-title">📦 材质参数</p>', unsafe_allow_html=True)
    material = st.selectbox("材质选择", list(MATERIAL_DB.keys()), index=0)

    if material == "自定义":
        fy = st.number_input("材料屈服强度 Fy (MPa)", min_value=100.0, max_value=1000.0, value=355.0, step=5.0)
        fu = st.number_input("材料抗拉强度 Fu (MPa)", min_value=100.0, max_value=1200.0, value=470.0, step=5.0)
    else:
        fy = float(st.number_input("材料屈服强度 Fy (MPa)", value=float(MATERIAL_DB[material]["fy"]), step=5.0))
        fu = float(st.number_input("材料抗拉强度 Fu (MPa)", value=float(MATERIAL_DB[material]["fu"]), step=5.0))

    ns = st.number_input("结构安全系数 Ns", min_value=1.0, max_value=10.0, value=3.0, step=0.1,
                         help="默认取 3.0，大型起重设备常用安全系数")

    st.divider()

    # —— 工况载荷参数 ——
    st.markdown('<p class="sidebar-section-title">⚡ 工况载荷</p>', unsafe_allow_html=True)
    fv_ton = st.number_input("单只吊耳受力 Fv (t)", min_value=0.01, max_value=5000.0, value=10.0, step=1.0,
                              help='输入质量单位"吨"，系统自动换算为 kN（1t ≈ 9.81 kN）')
    fv = fv_ton * 9.81  # 吨 → kN
    st.markdown(f'<p style="color:#63b3ed; font-size:0.8rem; margin-top:-0.5rem;">≈ <b>{fv:.2f} kN</b></p>',
                unsafe_allow_html=True)
    kd = st.number_input("动载系数 Kd", min_value=1.0, max_value=3.0, value=1.25, step=0.05,
                         help="静载取 1.0，动载通常取 1.25")
    theta = st.slider("吊索与水平面夹角 θ (°)", min_value=10, max_value=90, value=60, step=1,
                      help="夹角越大，垂直分力越大，通常60°~90°")

    st.divider()

    # —— 几何尺寸 ——
    st.markdown('<p class="sidebar-section-title">📐 吊耳几何尺寸 (mm)</p>', unsafe_allow_html=True)
    t  = st.number_input("吊耳主板厚度 t (mm)",      min_value=1.0, max_value=200.0, value=20.0, step=1.0)
    t_p = st.number_input("单侧加强板厚度 tp (mm)",  min_value=0.0, max_value=100.0, value=0.0, step=1.0,
                          help="无加强板时取 0。有效厚度 t_eff = t + 2×tp")
    if t_p > 0:
        st.info(f"有效厚度 t_eff = {t} + 2×{t_p} = **{t + 2*t_p:.0f} mm**")
    d  = st.number_input("销轴孔孔径 d (mm)",        min_value=1.0, max_value=500.0, value=52.0, step=1.0)
    dp = st.number_input("销轴直径 dp (mm)",         min_value=1.0, max_value=500.0, value=50.0, step=1.0,
                         help="销轴直径，通常略小于孔径")
    R  = st.number_input("孔心至边缘距离 R (mm)",    min_value=1.0, max_value=1000.0, value=75.0, step=1.0)
    hf = st.number_input("焊缝焊脚尺寸 hf (mm)",    min_value=1.0, max_value=50.0, value=12.0, step=1.0)

    use_custom_W = st.checkbox("自定义吊耳根部宽度 W", value=False,
                               help="不勾选则默认取 W = 2R（半圆头吊耳）")
    if use_custom_W:
        W = st.number_input("吊耳根部宽度 W (mm)", min_value=1.0, max_value=2000.0,
                            value=float(2 * R), step=1.0)
    else:
        W = 2 * R
        st.info(f"W = 2R = **{W:.0f} mm**")

    st.divider()
    st.markdown('<p style="color:#f7fafc; font-size:0.85rem; font-weight:500;">💡 公式参考：通用工程力学校验方法</p>', unsafe_allow_html=True)

# ─────────────────────────────────────────────
#  主界面 — 计算 & 输出
# ─────────────────────────────────────────────

# 页头
st.markdown('<h1 class="main-title">🔩 吊耳应力校验计算书</h1>', unsafe_allow_html=True)
st.markdown('<p class="subtitle">LIFTING LUG STRESS VERIFICATION REPORT</p>', unsafe_allow_html=True)

# 输入有效性检查
if dp > d:
    st.error("❌ 销轴直径 dp 不能大于孔径 d，请修正输入参数。")
    st.stop()

# 执行统一计算
results = calculate_lug_stresses(
    fv=fv, kd=kd, theta_deg=theta,
    t=t, t_p=t_p, d=d, dp=dp, r=R, w=W, hf=hf,
    fy=fy, fu=fu, ns=ns,
)

# 检查是否有计算错误
if "error" in results:
    st.error(f"❌ {results['error']}")
    st.stop()

# 提取结果
P = results["p_design"]
t_eff = results["t_eff"]
areas = results["areas"]
metrics = results["metrics"]

checks = [
    metrics["tensile"]["passed"],
    metrics["tear_out"]["passed"],
    metrics["bearing"]["passed"],
    metrics["weld"]["passed"],
]
all_pass = all(checks)

# ── 综合状态横幅 ──
if all_pass:
    st.success("✅ **综合校验：全部通过 PASS** — 该吊耳设计满足所有受力校验要求。")
else:
    fail_count = sum(1 for c in checks if not c)
    st.error(f"🚨 **综合校验：存在风险 FAIL** — 共 {fail_count} 项指标超出许用应力，请修改参数！")

st.markdown('<hr class="section-divider">', unsafe_allow_html=True)

# ── 吊耳几何图 + 设计拉力摘要（并排布局）──
col_diagram, col_metrics = st.columns([2, 1])

with col_diagram:
    st.markdown(
        '<div style="background:rgba(26,32,55,0.6); border:1px solid rgba(99,179,237,0.15); '
        'border-radius:12px; padding:0.8rem; text-align:center;">'
        '<p style="color:#63b3ed; font-size:0.85rem; font-weight:600; margin-bottom:0.3rem;">'
        '📐 吊耳几何示意图</p>'
        '<p style="color:#718096; font-size:0.7rem; margin-bottom:0;">'
        '图示随左侧参数实时更新</p></div>',
        unsafe_allow_html=True,
    )
    try:
        fig_preview = draw_lug_diagram(W, R, d, t, t_p, theta)
        st.pyplot(fig_preview, use_container_width=True)
        plt.close(fig_preview)
    except Exception as e:
        st.error(f"预览图生成失败: {e}")

with col_metrics:
    # 设计拉力摘要
    st.metric("设计拉力 P", f"{P/1000:.2f} kN", help="已含动载系数 Kd 和角度修正")
    st.metric("动载系数 Kd", f"{kd}")
    st.metric("吊索角度 θ", f"{theta}°")
    st.metric("安全系数 Ns", f"{ns}")
    if t_p > 0:
        st.metric("有效厚度 t_eff", f"{t_eff:.0f} mm", help=f"t + 2×tp = {t} + 2×{t_p}")

st.markdown('<hr class="section-divider">', unsafe_allow_html=True)

# ── 四项应力状态卡片 ──
st.subheader("📊 应力校验结果")

indicators = [
    {
        "name": "① 销轴孔净截面拉应力",
        "symbol": "σ_t",
        "actual": metrics["tensile"]["value"],
        "allowable": metrics["tensile"]["allowable"],
        "unit": "MPa",
        "pass": metrics["tensile"]["passed"],
    },
    {
        "name": "② 孔端面双面剪切撕裂应力",
        "symbol": "τ_v",
        "actual": metrics["tear_out"]["value"],
        "allowable": metrics["tear_out"]["allowable"],
        "unit": "MPa",
        "pass": metrics["tear_out"]["passed"],
    },
    {
        "name": "③ 销轴孔承压应力",
        "symbol": "σ_b",
        "actual": metrics["bearing"]["value"],
        "allowable": metrics["bearing"]["allowable"],
        "unit": "MPa",
        "pass": metrics["bearing"]["passed"],
    },
    {
        "name": "④ 底部焊缝综合剪应力",
        "symbol": "τ_weld",
        "actual": metrics["weld"]["value"],
        "allowable": metrics["weld"]["allowable"],
        "unit": "MPa",
        "pass": metrics["weld"]["passed"],
    },
]

cols = st.columns(2)
for i, ind in enumerate(indicators):
    pct = (ind["actual"] / ind["allowable"]) * 100
    badge = (
        '<span class="status-badge-pass">✓ PASS</span>'
        if ind["pass"]
        else '<span class="status-badge-fail">✗ FAIL</span>'
    )
    card_class = "status-pass" if ind["pass"] else "status-fail"
    bar_color  = "#48bb78" if ind["pass"] else "#fc8181"
    bar_pct    = min(pct, 100)

    html = f"""
    <div class="status-card {card_class}">
        <div class="stress-name">{ind['name']} &nbsp; {badge}</div>
        <div class="stress-values">
            <span class="stress-actual">{ind['actual']:.2f}</span>
            <span class="stress-allowable">/ {ind['allowable']:.2f} {ind['unit']}<br>({ind['symbol']})</span>
        </div>
        <div style="margin-top:0.6rem;">
            <div style="background:rgba(255,255,255,0.05); border-radius:4px; height:6px; overflow:hidden;">
                <div style="background:{bar_color}; width:{bar_pct:.1f}%; height:100%;
                            border-radius:4px; transition:width 0.5s;"></div>
            </div>
            <div style="font-size:0.72rem; color:#718096; margin-top:0.25rem;">
                利用率 {pct:.1f}%
            </div>
        </div>
    </div>
    """
    with cols[i % 2]:
        st.markdown(html, unsafe_allow_html=True)

st.markdown('<hr class="section-divider">', unsafe_allow_html=True)

# ── 详细计算书 ──
st.markdown("---")
col_title, col_btn = st.columns([0.8, 0.2])
with col_title:
    st.subheader("📋 详细计算书")
with col_btn:
    # 准备导出数据
    params_dict = {
        'fv_ton': fv_ton, 'fv': fv, 'kd': kd, 'theta': theta,
        't': t, 'tp': t_p, 'd': d, 'dp': dp, 'r': R, 'w': W,
        'hf': hf, 'fy': fy, 'fu': fu, 'ns': ns
    }
    if "error" not in results:
        word_file = generate_word_report(params_dict, results)
        st.download_button(
            label="📥 导出 Word 报告",
            data=word_file,
            file_name="吊耳校验计算书_Lifting_Lug_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            help="点击下载详细计算书至本地。由于网页应用特性，会调用浏览器的「另存为」对话框让您选择保存位置。"
        )

with st.expander("展开查看完整计算过程（可截图存档）", expanded=True):
    st.markdown(f"""
#### 一、基本参数
| 参数 | 符号 | 数值 | 单位 |
|------|------|------|------|
| 单只吊耳受力 | $F_v$ | {fv_ton} t（{fv:.2f} kN） | — |
| 动载系数 | $K_d$ | {kd} | — |
| 吊索夹角 | $θ$ | {theta} | ° |
| 材料屈服强度 | $F_y$ | {fy} | MPa |
| 材料抗拉强度 | $F_u$ | {fu} | MPa |
| 结构安全系数 | $N_s$ | {ns} | — |
| 吊耳主板厚 | $t$ | {t} | mm |
| 单侧加强板厚 | $t_p$ | {t_p} | mm |
| **有效厚度** | $t_{{eff}}$ | **{t_eff:.0f}** | mm |
| 销轴孔孔径 | $d$ | {d} | mm |
| 销轴直径 | $d_p$ | {dp} | mm |
| 孔心至边缘距离 | $R$ | {R} | mm |
| 角焊缝焊脚 | $h_f$ | {hf} | mm |
| 吊耳根部宽度 | $W$ | {W:.1f} | mm |
""")

    if t_p > 0:
        st.markdown(f"""
> **注：** 有效厚度 $t_{{eff}} = t + 2 \\times t_p = {t} + 2 \\times {t_p} = {t_eff:.0f}$ mm。应力计算（指标一~三）采用有效厚度，焊缝应力（指标四）仅通过主板传递。
""")

    st.markdown("---")
    st.markdown(f"""
#### 二、设计拉力换算

$$P = \\frac{{F_v \\times 1000 \\times K_d}}{{\\sin\\theta}} = \\frac{{{fv} \\times 1000 \\times {kd}}}{{\\sin({theta}°)}} = \\frac{{{fv*1000*kd:,.0f}}}{{{math.sin(math.radians(theta)):.4f}}} = \\mathbf{{{P:,.0f} \\ \\text{{N}}}} = \\mathbf{{{P/1000:.2f} \\ \\text{{kN}}}}$$
""")

    st.markdown("---")
    st.markdown(f"""
#### 三、指标一：销轴孔净截面拉应力 $\\sigma_t$

许用拉应力：${ALLOWABLE_RULES['tensile']['label']} = \\dfrac{{{fy}}}{{{ns}}} = {metrics['tensile']['allowable']:.2f} \\ \\text{{MPa}}$

净截面面积：$A_{{net}} = 2 \\times t_{{eff}} \\times (R - d/2) = 2 \\times {t_eff:.0f} \\times ({R} - {d}/2) = {areas['net_area']:.2f} \\ \\text{{mm}}^2$

$$\\sigma_t = \\frac{{P}}{{A_{{net}}}} = \\frac{{{P:,.0f}}}{{{areas['net_area']:.2f}}} = \\mathbf{{{metrics['tensile']['value']:.2f} \\ \\text{{MPa}}}}$$

校验：$\\sigma_t = {metrics['tensile']['value']:.2f} \\ \\text{{MPa}} \\ {"\\le" if checks[0] else ">"} \\ [\\sigma_t] = {metrics['tensile']['allowable']:.2f} \\ \\text{{MPa}}$ → {"**✅ 通过**" if checks[0] else "**❌ 不满足**"}
""")

    st.markdown("---")
    st.markdown(f"""
#### 四、指标二：孔端面双面剪切撕裂应力 $\\tau_v$

许用剪切应力：${ALLOWABLE_RULES['tear_out']['label']} = \\dfrac{{0.58 \\times {fy}}}{{{ns}}} = {metrics['tear_out']['allowable']:.2f} \\ \\text{{MPa}}$

双面剪切面积：$A_{{shear}} = 2 \\times t_{{eff}} \\times (R - d/2) = 2 \\times {t_eff:.0f} \\times ({R} - {d}/2) = {areas['shear_area']:.2f} \\ \\text{{mm}}^2$

$$\\tau_v = \\frac{{P}}{{A_{{shear}}}} = \\frac{{{P:,.0f}}}{{{areas['shear_area']:.2f}}} = \\mathbf{{{metrics['tear_out']['value']:.2f} \\ \\text{{MPa}}}}$$

校验：$\\tau_v = {metrics['tear_out']['value']:.2f} \\ \\text{{MPa}} \\ {"\\le" if checks[1] else ">"} \\ [\\tau_v] = {metrics['tear_out']['allowable']:.2f} \\ \\text{{MPa}}$ → {"**✅ 通过**" if checks[1] else "**❌ 不满足**"}
""")

    st.markdown("---")
    st.markdown(f"""
#### 五、指标三：销轴孔局部承压应力 $\\sigma_b$

许用承压应力：${ALLOWABLE_RULES['bearing']['label']} = \\dfrac{{1.25 \\times {fy}}}{{{ns}}} = {metrics['bearing']['allowable']:.2f} \\ \\text{{MPa}}$

承压接触面积：$A_{{bearing}} = t_{{eff}} \\times d_p = {t_eff:.0f} \\times {dp} = {areas['bearing_area']:.2f} \\ \\text{{mm}}^2$

$$\\sigma_b = \\frac{{P}}{{A_{{bearing}}}} = \\frac{{{P:,.0f}}}{{{areas['bearing_area']:.2f}}} = \\mathbf{{{metrics['bearing']['value']:.2f} \\ \\text{{MPa}}}}$$

校验：$\\sigma_b = {metrics['bearing']['value']:.2f} \\ \\text{{MPa}} \\ {"\\le" if checks[2] else ">"} \\ [\\sigma_b] = {metrics['bearing']['allowable']:.2f} \\ \\text{{MPa}}$ → {"**✅ 通过**" if checks[2] else "**❌ 不满足**"}
""")

    st.markdown("---")
    st.markdown(f"""
#### 六、指标四：底部焊缝综合剪应力 $\\tau_{{weld}}$

许用焊缝应力：${ALLOWABLE_RULES['weld']['label']} = 0.3 \\times {fu} = {metrics['weld']['allowable']:.2f} \\ \\text{{MPa}}$

焊喉有效厚度：$a = 0.707 \\times h_f = 0.707 \\times {hf} = {areas['weld_throat']:.3f} \\ \\text{{mm}}$

有效焊缝面积：$A_{{weld}} = 0.707 \\times h_f \\times 2W = {areas['weld_throat']:.3f} \\times 2 \\times {W:.1f} = {areas['weld_area']:.2f} \\ \\text{{mm}}^2$

$$\\tau_{{weld}} = \\frac{{P}}{{0.707 \\times h_f \\times 2W}} = \\frac{{{P:,.0f}}}{{{areas['weld_area']:.2f}}} = \\mathbf{{{metrics['weld']['value']:.2f} \\ \\text{{MPa}}}}$$

校验：$\\tau_{{weld}} = {metrics['weld']['value']:.2f} \\ \\text{{MPa}} \\ {"\\le" if checks[3] else ">"} \\ [\\tau_{{weld}}] = {metrics['weld']['allowable']:.2f} \\ \\text{{MPa}}$ → {"**✅ 通过**" if checks[3] else "**❌ 不满足**"}
""")

    st.markdown("---")

    # 汇总表
    summary_data = {
        "校验项目": ["① 拉应力 σ_t", "② 撕裂应力 τ_v", "③ 承压应力 σ_b", "④ 焊缝应力 τ_weld"],
        "实际值 (MPa)": [
            f"{metrics['tensile']['value']:.2f}",
            f"{metrics['tear_out']['value']:.2f}",
            f"{metrics['bearing']['value']:.2f}",
            f"{metrics['weld']['value']:.2f}",
        ],
        "许用值 (MPa)": [
            f"{metrics['tensile']['allowable']:.2f}",
            f"{metrics['tear_out']['allowable']:.2f}",
            f"{metrics['bearing']['allowable']:.2f}",
            f"{metrics['weld']['allowable']:.2f}",
        ],
        "利用率 (%)": [
            f"{metrics['tensile']['value']/metrics['tensile']['allowable']*100:.1f}%",
            f"{metrics['tear_out']['value']/metrics['tear_out']['allowable']*100:.1f}%",
            f"{metrics['bearing']['value']/metrics['bearing']['allowable']*100:.1f}%",
            f"{metrics['weld']['value']/metrics['weld']['allowable']*100:.1f}%",
        ],
        "结论": [
            "✅ 通过" if checks[0] else "❌ 不满足",
            "✅ 通过" if checks[1] else "❌ 不满足",
            "✅ 通过" if checks[2] else "❌ 不满足",
            "✅ 通过" if checks[3] else "❌ 不满足",
        ],
    }
    df = pd.DataFrame(summary_data)
    st.markdown("#### 七、汇总表")
    st.dataframe(df, use_container_width=True, hide_index=True)

# ── 页脚 ──
st.markdown('<hr class="section-divider">', unsafe_allow_html=True)
st.markdown(
    '<p style="text-align:center; color:#4a5568; font-size:0.75rem;">'
    '吊耳计算器 v1.1 &nbsp;|&nbsp; 支持加强板有效厚度 &nbsp;|&nbsp; 仅供工程参考，最终设计须经专业工程师审核'
    '</p>',
    unsafe_allow_html=True,
)
